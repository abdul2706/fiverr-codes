import os
import re
import pandas as pd
from docx import Document

columns = ['code', 'statement', 'option_a', 'option_b', 'option_c', 'option_d', 'answer', 'explanation']
docx_base_path = 'bullet-explanations'
docx_files = os.listdir(docx_base_path)
docx_files = list(filter(lambda x: x.endswith('docx'), docx_files))
excel_base_path = 'excel-files2'
os.makedirs(excel_base_path, exist_ok=True)
text_base_path = 'text-files2'
os.makedirs(text_base_path, exist_ok=True)
error_log_file = open('errors2.txt', 'w', encoding='utf-8')

for docx_file_name in docx_files:
    # if 'endo-generated-151412' not in docx_file_name: continue
    if '.~' in docx_file_name: continue
    docx_file_path = os.path.join(docx_base_path, docx_file_name)
    print('Processing File:', docx_file_path)
    excel_file_path = os.path.join(docx_base_path, docx_file_name[:-5] + '.xlsx')
    df_excel = pd.read_excel(excel_file_path, sheet_name=0)
    # print('[df_excel]\n', df_excel)
    # print(df_excel.columns)
    try:
        document = Document(docx_file_path)
    except:
        print(f'error reading file {docx_file_path}')
        continue
    # print('[document]', len(document.paragraphs))
    data_lines = []
    for para in document.paragraphs:
        text = para.text.strip()
        if len(text) > 0:
            data_lines.append(text)
            # text_file.write(text + '\n')
    # print('[data_lines]', len(data_lines))

    text_file = open(os.path.join(text_base_path, docx_file_name[:-5] + '.txt'), 'w')
    start_index, end_index = -1, -1
    # hack for processing last MCQ as well
    # without it last MCQ is missed due to boundary condition of an MCQ
    total_rows = len(data_lines)
    data_lines.append(data_lines[0])
    try:
        for i, data_line in enumerate(data_lines):
            if data_line.startswith('MCQ'):
                # data_line = data_line[len('MCQ Code:'):].strip()
                code = re.findall(r'\d+', data_line)
            else:
                code = re.findall(r'^\d+', data_line)
            if code:
                end_index = i

                if start_index >= 0 and end_index >= 0:
                    if i < total_rows: text_file.write('\n')
                    current_list = data_lines[start_index:end_index]
                    # print('-'*25)
                    # print('[prev_code]', prev_code)
                    # print(current_list)
                    # print(df_excel.loc[df_excel.code == prev_code].shape)
                    # print('-'*25)
                    if 'code' in df_excel:
                        df_excel.loc[df_excel['code'] == prev_code, 'explanation'] = '\n'.join(current_list[1:])
                    elif 'code.1' in df_excel:
                        df_excel.loc[df_excel['code.1'] == prev_code, 'explanation'] = '\n'.join(current_list[1:])
                    else:
                        raise Exception("excel sheet doesn't contain column 'code'")

                prev_code = int(code[0])
                start_index = end_index
            if i < total_rows: text_file.write(data_line + '\n')
            code = None
    except Exception as e:
        print(e)
        error_log_file.write(f'error while processing file: {docx_file_path}')
        error_log_file.write(f'current data_line: {data_line}')
        error_log_file.write(f'exception: {str(e)}')

    text_file.close()
    excel_file_path = os.path.join(excel_base_path, docx_file_name[:-5] + '.xlsx')
    print('Saving File:', excel_file_path, '; total rows ->', df_excel.shape)
    df_excel.to_excel(excel_file_path, index=False)
