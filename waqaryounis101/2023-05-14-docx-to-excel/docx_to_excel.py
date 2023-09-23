import os
import re
import pandas as pd
from docx import Document

debug = False
columns = ['code', 'statement', 'option_a', 'option_b', 'option_c', 'option_d', 'answer', 'explanation']
docx_base_path = 'docx-files'
docx_files = os.listdir(docx_base_path)
excel_base_path = 'excel_output-docx_to_excel'
os.makedirs(excel_base_path, exist_ok=True)
text_base_path = 'text_output-docx_to_excel'
os.makedirs(text_base_path, exist_ok=True)
error_log_file = open('errors-docx_to_excel.txt', 'w', encoding='utf-8')

for docx_file_name in docx_files:
    # if not (
        # 'CVS-GENERATED-SK151412' in docx_file_name
        # 'RESPO-GENERATED-SK151412' in docx_file_name
        # 'RESPO-GENERATED-SK151412-new' in docx_file_name
        # 'musculo-generated-151412' in docx_file_name
        # 'MICRO-GENERATED-151412' in docx_file_name
        # 'git-patho-generated-151412' in docx_file_name
        # 'hematology-generated-151412-new' in docx_file_name
    # ): continue
    docx_file_path = os.path.join(docx_base_path, docx_file_name)
    print('Processing File:', docx_file_path)
    try:
        document = Document(docx_file_path)
    except:
        print(f'error reading file {docx_file_path}')
        continue
    # if debug: print('[document]', len(document.paragraphs))
    data_lines = []
    for para in document.paragraphs:
        text = para.text.strip()
        # \d+\s*:
        if len(text) > 0:
        # and (
            # re.findall('^\d+', text) or
            # re.findall('^[Aa]\)', text) or
            # re.findall('^[Bb]\)', text) or
            # re.findall('^[Cc]\)', text) or
            # re.findall('^[Dd]\)', text) or
            # re.findall('^ANSWER', text) or
            # re.findall('^EXPLANATION', text)
        # ):
            data_lines += text.split('\n')
    # print('[data_lines]', len(data_lines))

    start_index, end_index = -1, -1
    text_file = open(os.path.join(text_base_path, docx_file_name[:-5] + '.txt'), 'w')
    for i, data_line in enumerate(data_lines):
        code = re.findall(r'^\d+', data_line)
        if code:
            end_index = i
            if start_index >= 0 and end_index >= 0:
                text_file.write('\n')
            start_index = end_index
        text_file.write(data_line + '\n')
    text_file.close()

    data_dict = {key: [] for key in columns}
    start_index, end_index, prev_code = -1, -1, -1
    # hack for processing last MCQ as well
    # without it last MCQ is missed due to boundary condition of an MCQ
    total_rows = len(data_lines)
    data_lines.append(data_lines[0])
    try:
        for i, data_line in enumerate(data_lines):
            code = re.findall(r'^\d+', data_line)
            # print(start_index, end_index, '[code, prev_code]', code, prev_code, data_line)
            if code:
                end_index = i

                if start_index > -1 and end_index > -1 and end_index - start_index > 5: # and end_index - start_index < 9:
                    current_list = data_lines[start_index:end_index]
                    current_list = [re.sub('\s', ' ', list_item) for list_item in current_list]
                    # print(start_index, end_index, current_list)
                    # if 'BNorepinephrine' in current_list[5]:
                    #     print(current_list)

                    # handle the case when option_a is in the first line, split it into two lines
                    search_option_a = re.search('\W[Aa][).]', current_list[0])
                    if search_option_a:
                        # print(search_option_a.start())
                        option_a_text = current_list[0][search_option_a.start()+1:]
                        current_list[0] = current_list[0][:search_option_a.start()+1]
                        current_list.insert(1, option_a_text)
                        # print(start_index, end_index, current_list)

                    if len(re.findall('^[Aa][).]', current_list[1])) == 0:
                        current_list[0] = current_list[0] + current_list[1]
                        current_list.pop(1)

                    if not (re.findall('[Aa][).]', current_list[1]) and 
                            re.findall('[Bb][).]', current_list[2]) and 
                            re.findall('[Cc][).]', current_list[3]) and 
                            re.findall('[Dd][).]', current_list[4]) and
                            re.findall('answer', current_list[5].lower())):
                        error_log_file.write(f'wrong format in file: {docx_file_path}' + '\n')
                        error_log_file.write('the order of options or answer may be wrong' + '\n')
                        error_log_file.write(f'start_index={start_index}; end_index={end_index}' + '\n')
                        error_log_file.write('\n'.join(current_list) + '\n\n')
                    else:
                        statement = re.sub(r'^\d+[\W]+', '', current_list[0])
                        option_a = current_list[1][current_list[1].index(' '):].strip()
                        option_b = current_list[2][current_list[2].index(' '):].strip()
                        option_c = current_list[3][current_list[3].index(' '):].strip()
                        option_d = current_list[4][current_list[4].index(' '):].strip()
                        options_map = {'A': option_a, 'B': option_b, 'C': option_c, 'D': option_d}
                        # answer = current_list[5][current_list[5].index(' '):].strip()
                        answer_line = current_list[5]
                        if debug: print('[answer_line]', answer_line)
                        answer_line = re.sub('ANSWER:', '', answer_line).strip()
                        if debug: print('[answer_line]', answer_line)
                        answer, explanation = None, None
                        if len(current_list) > 6:
                            current_list[6] = ' '.join(current_list[6:])
                            explanation = current_list[6][current_list[6].index(' '):].strip()

                        if len(answer_line) == 0:
                            error_log_file.write(f'wrong format in file: {docx_file_path}' + '\n')
                            error_log_file.write('answer not found' + '\n')
                            error_log_file.write(f'start_index={start_index}; end_index={end_index}' + '\n')
                            error_log_file.write('\n'.join(current_list) + '\n\n')
                        else:
                            if len(answer_line) < 3:
                                option_find = [answer_line[0]]
                            else:
                                option_find = re.findall('^[ABCDabcd]\W+', answer_line)
                            if debug: print('[option_find]', option_find)
                            if len(option_find) > 0:
                                option_char = option_find[0][0].upper()
                                if debug: print('[option_char]', option_char)
                                answer = options_map[option_char]
                                if debug: print('[answer]', answer)
                                answer_line = answer_line[len(option_find[0] + answer) + 1:]
                            if debug: print('[answer_line]', answer_line)
                            if answer is None:
                                next_check = True
                                answer_line_lower = answer_line.lower()
                                # check for exact equality
                                if option_a.lower() == answer_line_lower:
                                    answer = option_a
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[==][option_a]', answer)
                                elif option_b.lower() == answer_line_lower:
                                    answer = option_b
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[==][option_b]', answer)
                                elif option_c.lower() == answer_line_lower:
                                    answer = option_c
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[==][option_c]', answer)
                                elif option_d.lower() == answer_line_lower:
                                    answer = option_d
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[==][option_d]', answer)
                                if debug: print('[==][answer_line]', answer_line, next_check)

                                # check for answer in option
                                if next_check and answer_line_lower in option_a.lower():
                                    answer = option_a
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[answer in option][option_a]', answer)
                                elif next_check and answer_line_lower in option_b.lower():
                                    answer = option_b
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[answer in option][option_b]', answer)
                                elif next_check and answer_line_lower in option_c.lower():
                                    answer = option_c
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[answer in option][option_c]', answer)
                                elif next_check and answer_line_lower in option_d.lower():
                                    answer = option_d
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[answer in option][option_d]', answer)
                                if debug: print('[answer in option][answer_line]', answer_line, next_check)

                                # check for option in answer
                                if next_check and option_a.lower() in answer_line_lower:
                                    answer = option_a
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[option in answer][option_a]', answer)
                                elif next_check and option_b.lower() in answer_line_lower:
                                    answer = option_b
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[option in answer][option_b]', answer)
                                elif next_check and option_c.lower() in answer_line_lower:
                                    answer = option_c
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[option in answer][option_c]', answer)
                                elif next_check and option_d.lower() in answer_line_lower:
                                    answer = option_d
                                    answer_line = re.sub('^\W+', '', answer_line[len(answer):])
                                    next_check = False
                                    if debug: print('[option in answer][option_d]', answer)
                                if debug: print('[option in answer][answer_line]', answer_line, next_check)

                            if explanation is None and len(answer_line) > 0:
                                explanation = answer_line
                            if debug: print('[explanation]', explanation)

                            data_dict['code'].append(prev_code.strip())
                            data_dict['statement'].append(statement.strip())
                            data_dict['option_a'].append(option_a)
                            data_dict['option_b'].append(option_b)
                            data_dict['option_c'].append(option_c)
                            data_dict['option_d'].append(option_d)
                            data_dict['answer'].append(answer)
                            data_dict['explanation'].append(explanation)
                            # print('-'*25)
                else:
                    if end_index > 0 and start_index > 0:
                        error_log_file.write(f'wrong format in file: {docx_file_path}' + '\n')
                        error_log_file.write('number of lines per MCQ should be more then 5 and less then 9' + '\n')
                        error_log_file.write(f'start_index={start_index}; end_index={end_index}' + '\n')
                        error_log_file.write('\n'.join(data_lines[start_index:end_index]) + '\n\n')

                prev_code = code[0] if isinstance(code, list) else code
                start_index = end_index
    except Exception as e:
        error_string = f'error while processing file: {docx_file_path}' + '\n' + \
                       f'exception: {repr(e)}' + '\n' + \
                       f'start_index={start_index}; end_index={end_index}' + '\n' + \
                       '\n'.join(data_lines[start_index:end_index]) + '\n'
        # print(error_string)
        error_log_file.write(error_string)
        error_log_file.write('\n')

    text_file.close()
    excel_data = pd.DataFrame(data_dict)
    excel_file_path = os.path.join(excel_base_path, docx_file_name[:-5] + '.xlsx')
    print('Saving File:', excel_file_path, '; total rows ->', excel_data.shape)
    excel_data.to_excel(excel_file_path, index=False)

error_log_file.close()
