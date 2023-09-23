import os
import re
import pandas as pd
from docx import Document

excel_base_path = 'excel_output-docx_to_excel'
excel_files = os.listdir(excel_base_path)
os.makedirs(excel_base_path, exist_ok=True)
docx_base_path = 'docx_output-excel_to_docx'
os.makedirs(docx_base_path, exist_ok=True)
text_base_path = 'text_output-excel_to_docx'
os.makedirs(text_base_path, exist_ok=True)

for excel_file_name in excel_files:
    excel_file_path = os.path.join(excel_base_path, excel_file_name)
    print('Processing File:', excel_file_path)
    df_excel = pd.read_excel(excel_file_path)
    text_file = open(os.path.join(text_base_path, excel_file_name[:-5] + '.txt'), 'w')
    for idx, row in df_excel.iterrows():
        text_file.write(f"{idx + 1}. {row['statement']}\n")
        text_file.write(f"A) {row['option_a']}\n")
        text_file.write(f"B) {row['option_b']}\n")
        text_file.write(f"C) {row['option_c']}\n")
        text_file.write(f"D) {row['option_d']}\n")
        text_file.write(f"ANSWER: {row['answer']}\n")
        text_file.write(f"EXPLANATION: {row['explanation']} (code: {row['code']})\n\n")
    text_file.close()

    document = Document()
    for idx, row in df_excel.iterrows():
        document.add_paragraph(f"{idx + 1}. {row['statement']}")
        document.add_paragraph(f"A) {row['option_a']}")
        document.add_paragraph(f"B) {row['option_b']}")
        document.add_paragraph(f"C) {row['option_c']}")
        document.add_paragraph(f"D) {row['option_d']}")
        document.add_paragraph(f"ANSWER: {row['answer']}")
        document.add_paragraph(f"EXPLANATION: {row['explanation']} (code: {row['code']})")
        document.add_paragraph()
    document.save(os.path.join(docx_base_path, excel_file_name[:-5] + '.docx'))
